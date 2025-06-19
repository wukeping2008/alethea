"""
HTML转PDF和Word转换器
将Alethea平台三层架构说明图HTML文件转换为PDF和Word格式
"""

import os
import sys
import time
from pathlib import Path

def install_required_packages():
    """安装必需的包"""
    packages = [
        'pdfkit',
        'python-docx',
        'beautifulsoup4',
        'selenium',
        'webdriver-manager'
    ]
    
    for package in packages:
        try:
            __import__(package.replace('-', '_'))
            print(f"✓ {package} 已安装")
        except ImportError:
            print(f"正在安装 {package}...")
            os.system(f"pip install {package}")

def html_to_pdf_selenium(html_file, output_pdf):
    """使用Selenium将HTML转换为PDF"""
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.chrome.service import Service
        from webdriver_manager.chrome import ChromeDriverManager
        import base64
        
        print("使用Selenium Chrome驱动转换PDF...")
        
        # 设置Chrome选项
        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument('--window-size=1200,800')
        chrome_options.add_argument('--print-to-pdf')
        
        # 自动下载并设置ChromeDriver
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=chrome_options)
        
        # 获取HTML文件的绝对路径
        html_path = os.path.abspath(html_file)
        file_url = f"file:///{html_path.replace(os.sep, '/')}"
        
        print(f"正在加载HTML文件: {file_url}")
        driver.get(file_url)
        
        # 等待页面加载完成
        time.sleep(3)
        
        # 使用Chrome的打印功能生成PDF
        pdf_options = {
            'landscape': False,
            'displayHeaderFooter': False,
            'printBackground': True,
            'preferCSSPageSize': True,
            'paperWidth': 8.27,
            'paperHeight': 11.7,
            'marginTop': 0.4,
            'marginBottom': 0.4,
            'marginLeft': 0.4,
            'marginRight': 0.4
        }
        
        # 执行打印命令
        result = driver.execute_cdp_cmd('Page.printToPDF', pdf_options)
        
        # 保存PDF文件
        with open(output_pdf, 'wb') as f:
            f.write(base64.b64decode(result['data']))
        
        driver.quit()
        print(f"✓ PDF文件已生成: {output_pdf}")
        return True
        
    except Exception as e:
        print(f"Selenium转换失败: {e}")
        return False

def html_to_pdf_pdfkit(html_file, output_pdf):
    """使用pdfkit将HTML转换为PDF"""
    try:
        import pdfkit
        
        print("使用pdfkit转换PDF...")
        
        # 配置选项
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None,
            'print-media-type': None
        }
        
        # 转换HTML到PDF
        pdfkit.from_file(html_file, output_pdf, options=options)
        print(f"✓ PDF文件已生成: {output_pdf}")
        return True
        
    except Exception as e:
        print(f"pdfkit转换失败: {e}")
        print("提示: 可能需要安装wkhtmltopdf")
        return False

def html_to_word(html_file, output_docx):
    """将HTML内容转换为Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from bs4 import BeautifulSoup
        
        print("正在转换为Word文档...")
        
        # 读取HTML文件
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # 解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 创建Word文档
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 添加标题
        title = doc.add_heading('Alethea智能教育平台三层架构', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加副标题
        subtitle = doc.add_paragraph('构建新型人工智能驱动的高等教育学习环境')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(14)
        subtitle_format.italic = True
        
        doc.add_paragraph()  # 空行
        
        # 应用层
        doc.add_heading('🎓 应用层 - 智能化在线教育交互平台', 1)
        
        app_features = [
            ('🔬 在线仿真实验', [
                '电路仿真器 (CircuitJS集成)',
                '物理仿真 (PhET平台)',
                '数学可视化 (Desmos/GeoGebra)',
                '化学分子建模 (MolView)',
                '控制系统仿真 (Simulink)'
            ]),
            ('📚 项目制学习', [
                '智能小车项目',
                '人脸识别系统',
                '智能家居IoT',
                'PLC工业控制',
                'AI算法实现'
            ]),
            ('🎯 课程导入系统', [
                '多格式文档解析 (PDF/Word/PPT)',
                '智能内容分类',
                '知识点自动提取',
                '个人知识库构建',
                '学习路径推荐'
            ]),
            ('📊 学习分析', [
                '数字画像生成',
                '学习行为追踪',
                '知识掌握评估',
                '个性化推荐',
                '实时学习分析'
            ])
        ]
        
        for feature_title, feature_items in app_features:
            doc.add_heading(feature_title, 2)
            for item in feature_items:
                p = doc.add_paragraph(f'• {item}')
                p.style = 'List Bullet'
        
        # 特色功能
        highlight = doc.add_paragraph()
        highlight.add_run('🌟 Alethea特色功能: ').bold = True
        highlight.add_run('集成多学科仿真平台，支持电子、物理、数学、化学等领域的在线实验，结合项目制学习模式，提供从理论到实践的完整学习闭环。')
        
        doc.add_page_break()
        
        # 优化层
        doc.add_heading('⚡ 智能优化层 - AI内容优化与个性化引擎', 1)
        
        opt_features = [
            ('🧠 专业Prompt工程', [
                '学科专业化提示词模板',
                '上下文感知提示优化',
                '多轮对话状态管理',
                '用户画像驱动个性化',
                '实验内容智能生成'
            ]),
            ('📖 知识库增强', [
                '个人文档智能解析',
                '知识图谱构建',
                '语义相似度匹配',
                '上下文检索增强 (RAG)',
                '知识点关联分析'
            ]),
            ('🎯 内容质量控制', [
                'AI生成内容验证',
                '质量评分算法',
                '多媒体内容增强',
                '实验可行性检查',
                '学科准确性验证'
            ]),
            ('⚡ 性能优化', [
                '智能缓存策略',
                '响应时间优化',
                '负载均衡调度',
                '用户行为分析',
                '系统性能监控'
            ])
        ]
        
        for feature_title, feature_items in opt_features:
            doc.add_heading(feature_title, 2)
            for item in feature_items:
                p = doc.add_paragraph(f'• {item}')
                p.style = 'List Bullet'
        
        # 核心技术
        highlight = doc.add_paragraph()
        highlight.add_run('🔧 核心优化技术: ').bold = True
        highlight.add_run('结合用户个人知识库和专业提示词工程，实现AI回答的精准化和个性化，通过多层质量控制确保内容的专业性和准确性。')
        
        # 技术标签
        tech_tags = ['RAG检索增强', 'Prompt Engineering', '知识图谱', '语义分析', '质量评估', '缓存优化']
        tech_p = doc.add_paragraph('技术栈: ')
        for i, tag in enumerate(tech_tags):
            if i > 0:
                tech_p.add_run(' • ')
            run = tech_p.add_run(tag)
            run.bold = True
        
        doc.add_page_break()
        
        # 模型层
        doc.add_heading('🤖 AI模型层 - 多模型融合的智能底层架构', 1)
        
        model_features = [
            ('🌐 云端AI服务', [
                'Google Gemini (主力模型)',
                'Anthropic Claude (推理专家)',
                'OpenAI GPT系列',
                '阿里云通义千问Plus',
                '火山引擎DeepSeek'
            ]),
            ('🏠 本地部署', [
                'Ollama DeepSeek R1 (本地推理)',
                '离线模式支持',
                '数据隐私保护',
                '低延迟响应',
                '成本控制优化'
            ]),
            ('🧮 智能调度', [
                '问题类型自动识别',
                '模型能力匹配算法',
                '负载均衡策略',
                '故障自动切换',
                '成本效益优化'
            ]),
            ('🔄 备用机制', [
                '多级备用策略',
                '服务健康检测',
                '自动降级处理',
                '错误恢复机制',
                '服务可用性保障'
            ])
        ]
        
        for feature_title, feature_items in model_features:
            doc.add_heading(feature_title, 2)
            for item in feature_items:
                p = doc.add_paragraph(f'• {item}')
                p.style = 'List Bullet'
        
        # 模型选择策略
        highlight = doc.add_paragraph()
        highlight.add_run('🎯 模型选择策略: ').bold = True
        highlight.add_run('基于问题内容智能选择最适合的AI模型：编程问题优选DeepSeek，物理化学问题使用Claude，数学计算选择Gemini，确保每个领域都有专业的AI支持。')
        
        # 技术标签
        model_tags = ['Gemini 1.5 Flash', 'Claude 3 Sonnet', 'DeepSeek R1', '通义千问Plus', 'Ollama本地部署', '智能路由']
        model_p = doc.add_paragraph('模型技术: ')
        for i, tag in enumerate(model_tags):
            if i > 0:
                model_p.add_run(' • ')
            run = model_p.add_run(tag)
            run.bold = True
        
        doc.add_page_break()
        
        # 核心优势
        doc.add_heading('🚀 Alethea平台核心优势', 1)
        
        advantages = [
            ('多学科融合', '支持电子、物理、数学、化学、计算机等多个理工科领域'),
            ('实验导向', '集成第三方仿真平台，提供真实的在线实验体验'),
            ('AI驱动', '多模型智能调度，确保专业领域问题的精准回答'),
            ('个性化学习', '基于用户知识库和学习行为的智能推荐系统')
        ]
        
        for title, desc in advantages:
            p = doc.add_paragraph()
            p.add_run(f'• {title}: ').bold = True
            p.add_run(desc)
        
        # 保存Word文档
        doc.save(output_docx)
        print(f"✓ Word文档已生成: {output_docx}")
        return True
        
    except Exception as e:
        print(f"Word转换失败: {e}")
        return False

def main():
    """主函数"""
    print("=== Alethea平台架构图转换工具 ===\n")
    
    # 检查HTML文件是否存在
    html_file = "Alethea平台三层架构说明图.html"
    if not os.path.exists(html_file):
        print(f"错误: 找不到HTML文件 {html_file}")
        return
    
    print("正在安装必需的包...")
    install_required_packages()
    print()
    
    # 生成输出文件名
    base_name = "Alethea平台三层架构说明图"
    pdf_file = f"{base_name}.pdf"
    docx_file = f"{base_name}.docx"
    
    print("开始转换...")
    
    # 转换为PDF
    print("\n1. 转换为PDF...")
    pdf_success = False
    
    # 首先尝试使用Selenium
    if html_to_pdf_selenium(html_file, pdf_file):
        pdf_success = True
    else:
        # 如果Selenium失败，尝试pdfkit
        print("尝试使用pdfkit...")
        if html_to_pdf_pdfkit(html_file, pdf_file):
            pdf_success = True
    
    if not pdf_success:
        print("❌ PDF转换失败")
    
    # 转换为Word
    print("\n2. 转换为Word...")
    word_success = html_to_word(html_file, docx_file)
    
    if not word_success:
        print("❌ Word转换失败")
    
    # 总结
    print("\n=== 转换完成 ===")
    if pdf_success:
        print(f"✓ PDF文件: {pdf_file}")
    if word_success:
        print(f"✓ Word文件: {docx_file}")
    
    if pdf_success or word_success:
        print("\n转换成功！")
    else:
        print("\n转换失败，请检查错误信息。")

if __name__ == "__main__":
    main()
