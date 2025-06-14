"""
个人知识库路由
Personal Knowledge Base Routes
"""

from flask import Blueprint, request, jsonify, session
import os
import json
from datetime import datetime
from werkzeug.utils import secure_filename
import hashlib
import uuid
from models.user import db, User
from models.llm_models import llm_manager
import PyPDF2
import docx
import pdfplumber

# 可选导入PyMuPDF
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("PyMuPDF未安装，将使用其他PDF处理方法")

personal_knowledge_bp = Blueprint('personal_knowledge', __name__, url_prefix='/api/personal-knowledge')

# 允许的文件类型
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'md', 'ppt', 'pptx', 'xls', 'xlsx'}
UPLOAD_FOLDER = 'uploads/personal_docs'

def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def ensure_upload_folder():
    """确保上传文件夹存在"""
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

def extract_text_from_pdf(file_path):
    """从PDF文件提取文本 - 使用多种方法确保成功"""
    text = ""
    
    # 方法1: 使用pdfplumber (最准确)
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            print(f"PDF文本提取成功 (pdfplumber): {len(text)} 字符")
            return text.strip()
    except Exception as e:
        print(f"pdfplumber提取失败: {e}")
    
    # 方法2: 使用PyMuPDF (fitz) - 如果可用
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(file_path)
            text = ""
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n"
            doc.close()
            if text.strip():
                print(f"PDF文本提取成功 (PyMuPDF): {len(text)} 字符")
                return text.strip()
        except Exception as e:
            print(f"PyMuPDF提取失败: {e}")
    
    # 方法3: 使用PyPDF2 (备用)
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            print(f"PDF文本提取成功 (PyPDF2): {len(text)} 字符")
            return text.strip()
    except Exception as e:
        print(f"PyPDF2提取失败: {e}")
    
    print(f"所有PDF文本提取方法都失败了")
    return ""

def extract_text_from_docx(file_path):
    """从Word文档提取文本"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Word文档文本提取失败: {e}")
        return ""

def extract_text_from_txt(file_path):
    """从文本文件提取内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='gbk') as file:
                return file.read().strip()
        except Exception as e:
            print(f"文本文件读取失败: {e}")
            return ""
    except Exception as e:
        print(f"文本文件读取失败: {e}")
        return ""

def generate_summary(content, max_length=200):
    """生成内容摘要"""
    if len(content) <= max_length:
        return content
    else:
        # 简单的摘要生成：取前几句话
        sentences = content.split('。')
        summary = ""
        for sentence in sentences:
            if len(summary + sentence + "。") <= max_length:
                summary += sentence + "。"
            else:
                break
        return summary or content[:max_length] + "..."

# 个人文档模型（简化版，存储在session中）
class PersonalDocument:
    def __init__(self, user_id, filename, content, category='personal', subject='general', tags=None, description=''):
        self.doc_id = str(uuid.uuid4())
        self.user_id = user_id
        self.filename = filename
        self.original_filename = filename
        self.content = content
        self.content_summary = generate_summary(content)
        self.category = category
        self.subject = subject
        self.file_size = len(content.encode('utf-8'))
        self.upload_time = datetime.now().isoformat()
        self.tags = tags or []
        self.description = description
        self.is_public = False
    
    def to_dict(self):
        return {
            'doc_id': self.doc_id,
            'user_id': self.user_id,
            'filename': self.filename,
            'original_filename': self.original_filename,
            'content': self.content,
            'content_summary': self.content_summary,
            'category': self.category,
            'subject': self.subject,
            'file_size': self.file_size,
            'upload_time': self.upload_time,
            'tags': self.tags,
            'description': self.description,
            'is_public': self.is_public
        }

def get_user_documents(user_id):
    """获取用户的文档列表"""
    documents_key = f'user_documents_{user_id}'
    return session.get(documents_key, [])

def save_user_documents(user_id, documents):
    """保存用户的文档列表"""
    documents_key = f'user_documents_{user_id}'
    session[documents_key] = documents

@personal_knowledge_bp.route('/upload', methods=['POST'])
def upload_document():
    """上传文档到个人知识库"""
    try:
        # 检查用户登录状态（简化版）
        user_id = session.get('user_id', 1)  # 暂时使用固定用户ID
        
        if 'file' not in request.files:
            return jsonify({
                "success": False,
                "error": "没有选择文件"
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                "success": False,
                "error": "没有选择文件"
            }), 400
        
        if not allowed_file(file.filename):
            return jsonify({
                "success": False,
                "error": f"不支持的文件类型，支持的类型：{', '.join(ALLOWED_EXTENSIONS)}"
            }), 400
        
        # 获取表单数据
        category = request.form.get('category', 'personal')
        subject = request.form.get('subject', 'general')
        tags_str = request.form.get('tags', '')
        description = request.form.get('description', '')
        
        # 处理标签
        tags = [tag.strip() for tag in tags_str.split(',') if tag.strip()] if tags_str else []
        
        # 确保上传文件夹存在
        ensure_upload_folder()
        
        # 安全的文件名
        filename = secure_filename(file.filename)
        unique_filename = f"{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        # 保存文件
        file.save(file_path)
        
        # 提取文本内容
        content = ""
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.pdf':
            content = extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            content = extract_text_from_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            content = extract_text_from_txt(file_path)
        else:
            # 对于其他文件类型，暂时只保存文件信息
            content = f"文件类型：{file_ext}\n文件名：{filename}\n描述：{description}"
        
        if not content:
            os.remove(file_path)  # 删除无法提取内容的文件
            return jsonify({
                "success": False,
                "error": "无法提取文件内容"
            }), 400
        
        # 创建文档对象
        document = PersonalDocument(
            user_id=user_id,
            filename=unique_filename,
            content=content,
            category=category,
            subject=subject,
            tags=tags,
            description=description
        )
        document.original_filename = filename
        
        # 保存到session（实际应用中应该保存到数据库）
        user_documents = get_user_documents(user_id)
        user_documents.append(document.to_dict())
        save_user_documents(user_id, user_documents)
        
        return jsonify({
            "success": True,
            "document": document.to_dict(),
            "message": f"文档 {filename} 上传成功"
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"上传失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/documents')
def get_documents():
    """获取用户的文档列表"""
    try:
        user_id = session.get('user_id', 1)
        
        # 获取查询参数
        category = request.args.get('category', '')
        subject = request.args.get('subject', '')
        search = request.args.get('search', '')
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 10))
        
        # 获取用户文档
        user_documents = get_user_documents(user_id)
        
        # 筛选文档
        filtered_docs = []
        for doc_dict in user_documents:
            # 分类筛选
            if category and doc_dict.get('category') != category:
                continue
            
            # 学科筛选
            if subject and doc_dict.get('subject') != subject:
                continue
            
            # 搜索筛选
            if search:
                search_lower = search.lower()
                if (search_lower not in doc_dict.get('original_filename', '').lower() and
                    search_lower not in doc_dict.get('content', '').lower() and
                    search_lower not in doc_dict.get('description', '').lower()):
                    continue
            
            filtered_docs.append(doc_dict)
        
        # 分页
        total = len(filtered_docs)
        start = (page - 1) * per_page
        end = start + per_page
        documents = filtered_docs[start:end]
        
        return jsonify({
            "success": True,
            "documents": documents,
            "pagination": {
                "page": page,
                "per_page": per_page,
                "total": total,
                "pages": (total + per_page - 1) // per_page
            }
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"获取文档列表失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/search')
def search_documents():
    """搜索用户的文档"""
    try:
        user_id = session.get('user_id', 1)
        query = request.args.get('q', '')
        
        if not query:
            return jsonify({
                "success": False,
                "error": "搜索关键词不能为空"
            })
        
        # 获取用户文档
        user_documents = get_user_documents(user_id)
        
        # 搜索文档
        results = []
        query_lower = query.lower()
        
        for doc_dict in user_documents:
            # 计算相关性分数
            score = 0
            
            # 文件名匹配
            if query_lower in doc_dict.get('original_filename', '').lower():
                score += 10
            
            # 内容匹配
            content_lower = doc_dict.get('content', '').lower()
            score += content_lower.count(query_lower)
            
            # 描述匹配
            if query_lower in doc_dict.get('description', '').lower():
                score += 5
            
            # 标签匹配
            for tag in doc_dict.get('tags', []):
                if query_lower in tag.lower():
                    score += 8
            
            if score > 0:
                doc_dict['relevance_score'] = score
                results.append(doc_dict)
        
        # 按相关性排序
        results.sort(key=lambda x: x['relevance_score'], reverse=True)
        
        return jsonify({
            "success": True,
            "query": query,
            "results": results[:20],  # 限制返回20个结果
            "count": len(results)
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"搜索失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/chat', methods=['POST'])
def chat_with_knowledge():
    """基于个人知识库进行AI对话"""
    try:
        user_id = session.get('user_id', 1)
        data = request.get_json()
        message = data.get('message', '')
        history = data.get('history', [])
        
        if not message:
            return jsonify({
                "success": False,
                "error": "消息不能为空"
            })
        
        # 获取用户文档
        user_documents = get_user_documents(user_id)
        
        # 搜索相关文档
        relevant_docs = []
        message_lower = message.lower()
        
        for doc_dict in user_documents:
            score = 0
            
            # 简单的相关性计算
            content_lower = doc_dict.get('content', '').lower()
            filename_lower = doc_dict.get('original_filename', '').lower()
            
            # 计算关键词匹配
            words = message_lower.split()
            for word in words:
                if len(word) > 2:  # 忽略太短的词
                    score += content_lower.count(word) * 2
                    score += filename_lower.count(word) * 3
            
            if score > 0:
                doc_dict['relevance_score'] = score
                relevant_docs.append(doc_dict)
        
        # 按相关性排序，取前3个
        relevant_docs.sort(key=lambda x: x['relevance_score'], reverse=True)
        top_docs = relevant_docs[:3]
        
        # 构建上下文
        context_parts = []
        if top_docs:
            context_parts.append("基于你的个人知识库中的相关文档：")
            for i, doc in enumerate(top_docs, 1):
                content_preview = doc['content'][:500] + "..." if len(doc['content']) > 500 else doc['content']
                context_parts.append(f"\n文档{i}：{doc['original_filename']}")
                context_parts.append(f"内容摘要：{content_preview}")
        
        # 构建完整的提示词
        if context_parts:
            context_prompt = "\n".join(context_parts)
            full_prompt = f"""{context_prompt}

用户问题：{message}

请基于上述文档内容回答用户的问题。如果文档中没有直接相关的信息，请结合你的知识进行回答，并明确说明哪些信息来自文档，哪些是你的补充。回答要准确、有用且易于理解。"""
        else:
            full_prompt = f"""用户问题：{message}

你是用户的个人AI助手。虽然用户的知识库中暂时没有与此问题直接相关的文档，但请基于你的知识为用户提供有用的回答。建议用户上传相关文档以获得更个性化的帮助。"""
        
        # 调用AI模型
        import asyncio
        response = asyncio.run(llm_manager.generate_response(
            prompt=full_prompt,
            provider=None  # 自动选择
        ))
        
        if response.get('content'):
            return jsonify({
                "success": True,
                "response": response['content'],
                "related_documents": [
                    {
                        "doc_id": doc['doc_id'],
                        "filename": doc['original_filename'],
                        "relevance_score": doc['relevance_score']
                    } for doc in top_docs
                ],
                "provider_used": response.get('provider', 'unknown')
            })
        else:
            return jsonify({
                "success": False,
                "error": "AI回答生成失败"
            }), 500
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"对话失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/delete/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    """删除文档"""
    try:
        user_id = session.get('user_id', 1)
        
        # 获取用户文档
        user_documents = get_user_documents(user_id)
        
        # 查找要删除的文档
        doc_to_delete = None
        for doc_dict in user_documents:
            if doc_dict['doc_id'] == doc_id:
                doc_to_delete = doc_dict
                break
        
        if not doc_to_delete:
            return jsonify({
                "success": False,
                "error": "文档不存在"
            }), 404
        
        # 删除文件
        if 'filename' in doc_to_delete:
            file_path = os.path.join(UPLOAD_FOLDER, doc_to_delete['filename'])
            if os.path.exists(file_path):
                os.remove(file_path)
        
        # 从列表中移除
        user_documents = [doc for doc in user_documents if doc['doc_id'] != doc_id]
        save_user_documents(user_id, user_documents)
        
        return jsonify({
            "success": True,
            "message": "文档删除成功"
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"删除失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/stats')
def get_knowledge_stats():
    """获取个人知识库统计信息"""
    try:
        user_id = session.get('user_id', 1)
        
        # 获取用户文档
        user_documents = get_user_documents(user_id)
        
        # 统计信息
        total_docs = len(user_documents)
        total_content_size = sum(len(doc.get('content', '')) for doc in user_documents)
        
        # 按分类统计
        category_stats = {}
        subject_stats = {}
        
        for doc in user_documents:
            category = doc.get('category', 'other')
            subject = doc.get('subject', 'general')
            
            category_stats[category] = category_stats.get(category, 0) + 1
            subject_stats[subject] = subject_stats.get(subject, 0) + 1
        
        # 最近上传的文档
        recent_docs = sorted(user_documents, key=lambda x: x.get('upload_time', ''), reverse=True)[:5]
        
        return jsonify({
            "success": True,
            "stats": {
                "total_documents": total_docs,
                "total_content_size_kb": round(total_content_size / 1024, 2),
                "categories": category_stats,
                "subjects": subject_stats,
                "recent_uploads": [
                    {
                        "filename": doc['original_filename'],
                        "category": doc['category'],
                        "upload_time": doc['upload_time']
                    } for doc in recent_docs
                ]
            }
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"获取统计信息失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/document/<doc_id>')
def get_document(doc_id):
    """获取文档详情"""
    try:
        user_id = session.get('user_id', 1)
        
        # 获取用户文档
        user_documents = get_user_documents(user_id)
        
        # 查找文档
        for doc_dict in user_documents:
            if doc_dict['doc_id'] == doc_id:
                return jsonify({
                    "success": True,
                    "document": doc_dict
                })
        
        return jsonify({
            "success": False,
            "error": "文档不存在"
        }), 404
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"获取文档详情失败: {str(e)}"
        }), 500

# 分类管理相关API
@personal_knowledge_bp.route('/categories', methods=['GET'])
def get_categories():
    """获取用户的文档分类"""
    try:
        user_id = session.get('user_id', 1)
        user_documents = get_user_documents(user_id)
        
        # 统计分类信息
        categories = {}
        for doc in user_documents:
            category = doc.get('category', 'other')
            if category not in categories:
                categories[category] = {
                    'name': category,
                    'display_name': get_category_display_name(category),
                    'count': 0,
                    'description': get_category_description(category),
                    'color': get_category_color(category)
                }
            categories[category]['count'] += 1
        
        return jsonify({
            "success": True,
            "categories": list(categories.values())
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"获取分类失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/categories/<category_name>/documents', methods=['GET'])
def get_category_documents(category_name):
    """获取指定分类的文档"""
    try:
        user_id = session.get('user_id', 1)
        user_documents = get_user_documents(user_id)
        
        # 筛选指定分类的文档
        category_docs = [doc for doc in user_documents if doc.get('category') == category_name]
        
        return jsonify({
            "success": True,
            "category": category_name,
            "documents": category_docs,
            "count": len(category_docs)
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"获取分类文档失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/tags', methods=['GET'])
def get_tags():
    """获取所有标签"""
    try:
        user_id = session.get('user_id', 1)
        user_documents = get_user_documents(user_id)
        
        # 收集所有标签
        all_tags = {}
        for doc in user_documents:
            for tag in doc.get('tags', []):
                if tag not in all_tags:
                    all_tags[tag] = {
                        'name': tag,
                        'count': 0,
                        'documents': []
                    }
                all_tags[tag]['count'] += 1
                all_tags[tag]['documents'].append({
                    'doc_id': doc['doc_id'],
                    'filename': doc['original_filename']
                })
        
        return jsonify({
            "success": True,
            "tags": list(all_tags.values())
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"获取标签失败: {str(e)}"
        }), 500

# 个性化设置相关API
@personal_knowledge_bp.route('/settings', methods=['GET'])
def get_user_settings():
    """获取用户个性化设置"""
    try:
        user_id = session.get('user_id', 1)
        settings_key = f'user_settings_{user_id}'
        
        # 默认设置
        default_settings = {
            'ai_style': 'detailed',  # detailed, concise, academic, casual
            'knowledge_scope': 'all',  # all, category, recent
            'auto_categorize': True,
            'smart_tags': True,
            'language': 'zh-CN',
            'theme': 'light',
            'notifications': {
                'upload_success': True,
                'ai_response': True,
                'daily_summary': False
            },
            'privacy': {
                'default_visibility': 'private',
                'allow_sharing': False
            },
            'ai_preferences': {
                'preferred_provider': 'auto',
                'response_length': 'medium',
                'include_sources': True,
                'explain_reasoning': False
            }
        }
        
        # 获取用户设置，如果不存在则使用默认设置
        user_settings = session.get(settings_key, default_settings)
        
        return jsonify({
            "success": True,
            "settings": user_settings
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"获取设置失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/settings', methods=['POST'])
def update_user_settings():
    """更新用户个性化设置"""
    try:
        user_id = session.get('user_id', 1)
        settings_key = f'user_settings_{user_id}'
        
        data = request.get_json()
        if not data:
            return jsonify({
                "success": False,
                "error": "没有提供设置数据"
            }), 400
        
        # 获取当前设置
        current_settings = session.get(settings_key, {})
        
        # 更新设置
        current_settings.update(data)
        
        # 保存设置
        session[settings_key] = current_settings
        
        return jsonify({
            "success": True,
            "settings": current_settings,
            "message": "设置更新成功"
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"更新设置失败: {str(e)}"
        }), 500

@personal_knowledge_bp.route('/settings/reset', methods=['POST'])
def reset_user_settings():
    """重置用户设置为默认值"""
    try:
        user_id = session.get('user_id', 1)
        settings_key = f'user_settings_{user_id}'
        
        # 删除用户设置，下次获取时会使用默认值
        if settings_key in session:
            del session[settings_key]
        
        return jsonify({
            "success": True,
            "message": "设置已重置为默认值"
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"重置设置失败: {str(e)}"
        }), 500

# 辅助函数
def get_category_display_name(category):
    """获取分类的显示名称"""
    category_names = {
        'personal': '个人学习',
        'research': '研究资料',
        'course': '课程笔记',
        'reference': '参考文献',
        'project': '项目文档',
        'teaching': '教学资料',
        'other': '其他'
    }
    return category_names.get(category, category)

def get_category_description(category):
    """获取分类的描述"""
    descriptions = {
        'personal': '个人学习和知识积累相关的文档',
        'research': '学术研究、论文、实验数据等',
        'course': '课程学习笔记、作业、考试资料',
        'reference': '参考书籍、文献、资料汇编',
        'project': '项目相关的文档、报告、计划',
        'teaching': '教学用的课件、教案、习题',
        'other': '其他类型的文档'
    }
    return descriptions.get(category, '未分类的文档')

def get_category_color(category):
    """获取分类的颜色"""
    colors = {
        'personal': '#3b82f6',    # 蓝色
        'research': '#8b5cf6',    # 紫色
        'course': '#10b981',      # 绿色
        'reference': '#f59e0b',   # 橙色
        'project': '#ef4444',     # 红色
        'teaching': '#06b6d4',    # 青色
        'other': '#6b7280'        # 灰色
    }
    return colors.get(category, '#6b7280')
