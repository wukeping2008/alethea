#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
语料库文档转换器
将Alethea平台智能语料库系统详细介绍.md转换为Word格式
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_word_document(md_file_path, output_file_path):
    """将Markdown文件转换为Word文档"""
    
    # 创建新的Word文档
    doc = Document()
    
    # 设置文档样式
    setup_document_styles(doc)
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 解析并转换内容
    parse_markdown_to_word(content, doc)
    
    # 保存文档
    doc.save(output_file_path)
    print(f"Word文档已保存到: {output_file_path}")

def setup_document_styles(doc):
    """设置文档样式"""
    
    # 设置标题样式
    styles = doc.styles
    
    # 标题1样式
    if 'Heading 1' in styles:
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(20)
        heading1.font.bold = True
        heading1.font.name = '微软雅黑'
    
    # 标题2样式
    if 'Heading 2' in styles:
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(16)
        heading2.font.bold = True
        heading2.font.name = '微软雅黑'
    
    # 标题3样式
    if 'Heading 3' in styles:
        heading3 = styles['Heading 3']
        heading3.font.size = Pt(14)
        heading3.font.bold = True
        heading3.font.name = '微软雅黑'
    
    # 标题4样式
    if 'Heading 4' in styles:
        heading4 = styles['Heading 4']
        heading4.font.size = Pt(13)
        heading4.font.bold = True
        heading4.font.name = '微软雅黑'
    
    # 正文样式
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.font.size = Pt(11)
        normal.font.name = '微软雅黑'

def parse_markdown_to_word(content, doc):
    """解析Markdown内容并转换为Word格式"""
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # 空行
            i += 1
            continue
        
        elif line.startswith('# '):
            # 一级标题
            title = clean_emoji_text(line[2:].strip())
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('## '):
            # 二级标题
            title = clean_emoji_text(line[3:].strip())
            doc.add_heading(title, level=2)
            
        elif line.startswith('### '):
            # 三级标题
            title = clean_emoji_text(line[4:].strip())
            doc.add_heading(title, level=3)
            
        elif line.startswith('#### '):
            # 四级标题
            title = clean_emoji_text(line[5:].strip())
            doc.add_heading(title, level=4)
            
        elif line.startswith('- **') and line.endswith('**：'):
            # 带粗体的列表项（中文冒号）
            match = re.match(r'- \*\*(.*?)\*\*：\s*(.*)', line)
            if match:
                label, content = match.groups()
                p = doc.add_paragraph(style='List Bullet')
                run1 = p.add_run(label + '：')
                run1.bold = True
                p.add_run(content)
            else:
                doc.add_paragraph(clean_emoji_text(line[2:].strip()), style='List Bullet')
                
        elif line.startswith('- **') and ':' in line:
            # 带粗体的列表项（英文冒号）
            match = re.match(r'- \*\*(.*?)\*\*:\s*(.*)', line)
            if match:
                label, content = match.groups()
                p = doc.add_paragraph(style='List Bullet')
                run1 = p.add_run(label + '：')
                run1.bold = True
                p.add_run(content)
            else:
                doc.add_paragraph(clean_emoji_text(line[2:].strip()), style='List Bullet')
                
        elif line.startswith('- '):
            # 普通列表项
            doc.add_paragraph(clean_emoji_text(line[2:].strip()), style='List Bullet')
            
        elif line.startswith('1. ') or re.match(r'^\d+\. ', line):
            # 数字列表
            content = re.sub(r'^\d+\. ', '', line)
            doc.add_paragraph(clean_emoji_text(content), style='List Number')
            
        elif line.startswith('**') and line.endswith('**：'):
            # 粗体段落标题（中文冒号）
            label = line[2:-3].strip()
            p = doc.add_paragraph()
            run = p.add_run(clean_emoji_text(label) + '：')
            run.bold = True
            
        elif line.startswith('**') and line.endswith('**:'):
            # 粗体段落标题（英文冒号）
            label = line[2:-3].strip()
            p = doc.add_paragraph()
            run = p.add_run(clean_emoji_text(label) + '：')
            run.bold = True
            
        else:
            # 普通段落
            # 处理粗体文本和表情符号
            paragraph_text = process_text_formatting(line)
            if paragraph_text.strip():
                p = doc.add_paragraph()
                add_formatted_text(p, line)
        
        i += 1

def clean_emoji_text(text):
    """清理文本中的表情符号，保留中文内容"""
    # 移除常见的表情符号，但保留中文和英文内容
    emoji_pattern = r'[📚🔄🎯📖🏷️🧠🔍📊🚀💾🤖🌟🎓🔄📈📋✨🛡️⚡🇨🇳🔧🏗️📋]'
    cleaned = re.sub(emoji_pattern, '', text).strip()
    return cleaned

def process_text_formatting(text):
    """处理文本格式"""
    # 处理粗体标记
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # 清理表情符号
    text = clean_emoji_text(text)
    return text

def add_formatted_text(paragraph, text):
    """向段落添加格式化文本"""
    # 分割文本，处理粗体部分
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 粗体文本
            bold_text = clean_emoji_text(part[2:-2])
            if bold_text:
                run = paragraph.add_run(bold_text)
                run.bold = True
        else:
            # 普通文本
            normal_text = clean_emoji_text(part)
            if normal_text:
                paragraph.add_run(normal_text)

def main():
    """主函数"""
    md_file = "Alethea平台智能语料库系统详细介绍.md"
    word_file = "Alethea平台智能语料库系统详细介绍.docx"
    
    try:
        create_word_document(md_file, word_file)
        print("语料库文档转换完成！")
        print(f"输出文件：{word_file}")
    except FileNotFoundError:
        print(f"错误：找不到文件 {md_file}")
    except Exception as e:
        print(f"转换过程中出现错误：{e}")

if __name__ == "__main__":
    main()
