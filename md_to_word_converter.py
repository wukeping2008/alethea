#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Document Converter
将Markdown文档转换为Word格式
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_word_document(md_file_path, output_file_path):
    """将Markdown文件转换为Word文档"""
    
    # 创建新的Word文档
    doc = Document()
    
    # 设置文档样式
    setup_document_styles(doc)
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 解析并转换内容
    parse_markdown_to_word(content, doc)
    
    # 保存文档
    doc.save(output_file_path)
    print(f"Word文档已保存到: {output_file_path}")

def setup_document_styles(doc):
    """设置文档样式"""
    
    # 设置标题样式
    styles = doc.styles
    
    # 标题1样式
    if 'Heading 1' in styles:
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(18)
        heading1.font.bold = True
        heading1.font.name = '微软雅黑'
    
    # 标题2样式
    if 'Heading 2' in styles:
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(16)
        heading2.font.bold = True
        heading2.font.name = '微软雅黑'
    
    # 标题3样式
    if 'Heading 3' in styles:
        heading3 = styles['Heading 3']
        heading3.font.size = Pt(14)
        heading3.font.bold = True
        heading3.font.name = '微软雅黑'
    
    # 正文样式
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.font.size = Pt(12)
        normal.font.name = '微软雅黑'

def parse_markdown_to_word(content, doc):
    """解析Markdown内容并转换为Word格式"""
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # 空行
            i += 1
            continue
        
        elif line.startswith('# '):
            # 一级标题
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('## '):
            # 二级标题
            title = line[3:].strip()
            doc.add_heading(title, level=2)
            
        elif line.startswith('### '):
            # 三级标题
            title = line[4:].strip()
            doc.add_heading(title, level=3)
            
        elif line.startswith('**') and line.endswith('**:'):
            # 粗体标签（如镜头、操作等）
            label = line[2:-3].strip()
            p = doc.add_paragraph()
            run = p.add_run(label + ':')
            run.bold = True
            
        elif line.startswith('- **'):
            # 带粗体的列表项
            match = re.match(r'- \*\*(.*?)\*\*:\s*(.*)', line)
            if match:
                label, content = match.groups()
                p = doc.add_paragraph(style='List Bullet')
                run1 = p.add_run(label + ': ')
                run1.bold = True
                p.add_run(content)
            else:
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
                
        elif line.startswith('- '):
            # 普通列表项
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
            
        elif line.startswith('1. ') or re.match(r'^\d+\. ', line):
            # 数字列表
            content = re.sub(r'^\d+\. ', '', line)
            doc.add_paragraph(content, style='List Number')
            
        elif line.startswith('| '):
            # 表格处理
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # 回退一行
            
            if table_lines:
                create_table_from_markdown(doc, table_lines)
                
        elif line.startswith('```'):
            # 代码块
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                p = doc.add_paragraph()
                p.style = 'Normal'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                
        elif line.startswith('---'):
            # 分隔线
            doc.add_paragraph('─' * 50)
            
        else:
            # 普通段落
            # 处理粗体文本
            paragraph_text = process_bold_text(line)
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)
        
        i += 1

def process_bold_text(text):
    """处理文本中的粗体标记"""
    # 简单处理，移除markdown的粗体标记
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    return text

def create_table_from_markdown(doc, table_lines):
    """从Markdown表格创建Word表格"""
    
    # 解析表格数据
    rows = []
    for line in table_lines:
        if '---' in line:  # 跳过分隔行
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # 去掉首尾空元素
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    # 填充数据
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_data
                
                # 设置表头样式
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def main():
    """主函数"""
    md_file = "Alethea平台使用介绍录制剧本.md"
    word_file = "Alethea平台使用介绍录制剧本.docx"
    
    try:
        create_word_document(md_file, word_file)
        print("转换完成！")
    except FileNotFoundError:
        print(f"错误：找不到文件 {md_file}")
    except Exception as e:
        print(f"转换过程中出现错误：{e}")

if __name__ == "__main__":
    main()
